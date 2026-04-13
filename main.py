import ollama
import json
import re
from docx import Document

# --- 1. 系統設定區 --- (保持不變)
CONFIG = {
    "存證信函": {
        "template": "legal_notice.docx",
        "fields": ["sender_name", "sender_addr", "receiver_name", "receiver_addr", "fact_date", "amount", "deadline"]
    },
    "房屋租賃": {
        "template": "rent_template.docx",
        "fields": ["landlord", "tenant", "address", "start_date", "end_date", "rent", "pay_day", "special_terms"]
    }
}

FIELD_MAP = {
    "sender_name": "您的姓名", "sender_addr": "您的地址",
    "receiver_name": "債務人姓名", "receiver_addr": "債務人地址",
    "amount": "欠款金額", "fact_date": "借款日期", "deadline": "還款期限(天)",
    "landlord": "房東姓名", "tenant": "房客姓名", "address": "房屋地址",
    "start_date": "起租日期", "end_date": "退租日期", "rent": "每月租金",
    "pay_day": "每月交租日", "special_terms": "特別約定事項"
}

# --- 2. 邏輯處理區 --- (保持不變)
def generate_docx(template_name, data, output_name):
    try:
        doc = Document(template_name)
        all_paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)
        for para in all_paragraphs:
            for key, value in data.items():
                tag = f"{{{{{key}}}}}"
                if tag in para.text:
                    para.text = para.text.replace(tag, str(value))
        doc.save(output_name)
        print(f"\n✅ 文件已成功生成：{output_name}")
    except Exception as e:
        print(f"❌ 生成失敗: {e}")

# --- 3. Agent 對話核心：兩者兼具優化版 ---

def start_legal_agent():
    user_context = {}
    current_mode = None 
    
    print("⚖️ AI 法律助理已上線。您可以直接描述您的問題，或一次提供多項資訊。")
    
    while True:
        # A. 找出目前缺失的欄位
        next_field = None
        if current_mode:
            required = CONFIG[current_mode]['fields']
            missing = [f for f in required if f not in user_context or not str(user_context[f]).strip()]
            if missing:
                next_field = missing[0]

        user_input = input("\n👤 您：").strip()
        if not user_input: continue

        # B. 初始模式偵測
        if not current_mode:
            if any(k in user_input for k in ["欠", "錢", "還", "借"]): current_mode = "存證信函"
            elif any(k in user_input for k in ["租", "房", "合約"]): current_mode = "房屋租賃"

        # C. 【強力優化】全域掃描 + 問題意識 Prompt
        target_fields = CONFIG[current_mode]['fields'] if current_mode else "所有法律相關欄位"
        current_focus_label = FIELD_MAP.get(next_field, "初步情況")
        
        prompt = f"""
        你是一個精準的法律數據提取器。
        【目前背景】：使用者正在填寫「{current_mode if current_mode else "未知文件"}」。
        【正在詢問】：{current_focus_label} (對應欄位 Key: {next_field})

        ### 提取規則 ###
        1. **語言限制**：所有提取到的值必須使用【繁體中文】。
        2. **嚴禁翻譯**：絕對不要將內容翻譯成英文（例如：7天 不准變成 7 days）。
        3. **原樣提取**：保持使用者提供的單位與表達方式（例如：十萬元、2005年、7天）。
        4. 使用者現在的回覆極可能是為了回答「{current_focus_label}」，請優先提取並填入 "{next_field}"。
        5. 同時掃描輸入，看是否有其他資訊可填入：{target_fields}。
        6. 僅回傳 JSON，不要解釋。

        【輸出格式範例】：
        {{
            "extracted_data": {{
                "{next_field if next_field else 'info'}": "提取到的繁體中文值"
            }}
        }}

        使用者輸入："{user_input}"
        """
        
        try:
            response = ollama.chat(model='mistral', messages=[{'role': 'user', 'content': prompt}])
            raw_content = response['message']['content']
            
            # --- 除錯區：如果你又卡住了，把下面這行的註解拿掉，看 AI 回了什麼 ---
            # print(f"--- AI 回傳內容: {raw_content} ---") 

            match = re.search(r'\{.*\}', raw_content, re.DOTALL)
            if match:
                res = json.loads(match.group())
                
                # 更新數據
                new_data = res.get('extracted_data', {})
                for k, v in new_data.items():
                    # 只要值不是 null 或空的，就更新進去
                    if v and str(v).strip() not in ["null", "None", "未提供", "未知", ""]:
                        user_context[k] = v
            
            # D. 流程控制
            if not current_mode:
                print("🤖 AI：我不確定您的文件類型，請問是「存證信函」還是「房屋租賃」？")
                continue

            # 重新檢查缺失
            required = CONFIG[current_mode]['fields']
            missing = [f for f in required if f not in user_context]
            
            if not missing:
                print(f"🤖 AI：偵測到資訊已齊全！正在為您產出「{current_mode}」...")
                generate_docx(CONFIG[current_mode]['template'], user_context, f"Final_{current_mode}.docx")
                break
            else:
                next_f = missing[0]
                print(f"🤖 AI：已記錄。還缺「{FIELD_MAP.get(next_f, next_f)}」，請提供。")
                
        except Exception as e:
            print(f"⚠️ 處理中發生錯誤，請再試一次...")

if __name__ == "__main__":
    start_legal_agent()